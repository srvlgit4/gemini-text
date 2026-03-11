import time
import os
import threading
import telebot
from flask import Flask
from selenium import webdriver
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

# --- 1. FLASK DUMMY SERVER (For Render Free Tier) ---
app = Flask(__name__)

@app.route('/')
def health_check():
    return "Bot is alive and running!"

def run_server():
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)

# Start the web server in a separate background thread
threading.Thread(target=run_server, daemon=True).start()


# --- 2. TELEGRAM BOT SETUP ---
BOT_TOKEN = "8436841638:AAFz0JFN8fXxHqy5eQGFDLXeCUwn0JLcF4w"
bot = telebot.TeleBot(BOT_TOKEN)


# --- 3. THE MAGIC FORMATTING PARSER ---
def append_formatted_text(element, paragraph, is_user=False):
    for child in element.children:
        if child.name is None:
            text_str = str(child).replace('\n', ' ')
            if text_str.strip():
                run = paragraph.add_run(text_str)
                if is_user: 
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.bold = True
                    run.font.size = Pt(14)
        elif child.name in ['b', 'strong']:
            run = paragraph.add_run(child.get_text())
            run.bold = True
            if is_user: 
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(14)
        elif child.name in ['i', 'em']:
            run = paragraph.add_run(child.get_text())
            run.italic = True
            if is_user: 
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.bold = True
                run.font.size = Pt(14)
        else:
            append_formatted_text(child, paragraph, is_user)


# --- 4. SCRAPING FUNCTION ---
def extract_gemini_chat(url, output_filename):
    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-gpu') 

    # FIXED: Let Selenium 4's built-in manager handle the driver!
    driver = webdriver.Chrome(options=options)

    driver.get(url)
    time.sleep(6)
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    driver.quit()

    doc = Document()
    doc.add_heading('Extracted Gemini Chat', 0)

    last_text = ""
    search_tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'pre']

    for element in soup.find_all(search_tags):
        parent_tags = [p.name for p in element.parents]
        if element.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'] and any(tag in parent_tags for tag in ['table', 'ul', 'ol', 'pre']):
            continue
        if element.name in ['ul', 'ol'] and any(tag in parent_tags for tag in ['table', 'pre']):
            continue

        text = element.get_text(strip=True)
        if not text or text == last_text:
            continue
        last_text = text

        is_user_prompt = False
        for parent in element.parents:
            if parent.get('data-message-author-role') == 'user':
                is_user_prompt = True
                break
            class_str = " ".join(parent.get('class', [])).lower()
            if 'user' in class_str or 'query' in class_str:
                is_user_prompt = True
                break
            if parent.get('data-message-author-role') == 'model' or 'model' in class_str or 'response' in class_str:
                break

        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            if is_user_prompt:
                p = doc.add_paragraph()
                run = p.add_run("Question: ")
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(14)
            else:
                p = doc.add_heading(level=int(element.name[1]))
            append_formatted_text(element, p, is_user_prompt)
            
        elif element.name == 'p':
            p = doc.add_paragraph()
            if is_user_prompt:
                run = p.add_run("Question: ")
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                run.font.size = Pt(14)
            append_formatted_text(element, p, is_user_prompt)
            
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. ") 
                if is_user_prompt: 
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    run.bold = True
                    run.font.size = Pt(14)
                append_formatted_text(li, p, is_user_prompt)
                
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                append_formatted_text(li, p, is_user_prompt)
                
        elif element.name == 'table':
            rows = element.find_all('tr')
            if not rows: continue
            cols = rows[0].find_all(['td', 'th'])
            if not cols: continue
            
            table = doc.add_table(rows=len(rows), cols=len(cols))
            table.style = 'Table Grid'
            
            for r_idx, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for c_idx, cell in enumerate(cells):
                    if c_idx < len(cols): 
                        cell_paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        append_formatted_text(cell, cell_paragraph, is_user_prompt)
                        
        elif element.name == 'pre':
            p = doc.add_paragraph(text)
            p.style = 'Intense Quote' 

    doc.save(output_filename)


# --- 5. BOT COMMANDS ---
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.reply_to(message, "👋 Welcome! Send me a Google Gemini shared link, and I will format it into a clean DOCX file for you.")

# FIXED: Now checks for both standard and shortened (g.co) links
@bot.message_handler(func=lambda message: "gemini.google.com/share/" in message.text or "g.co/gemini/share/" in message.text)
def process_link(message):
    url = message.text.strip()
    processing_msg = bot.reply_to(message, "⏳ Scraping the chat and building your document... Please wait about 15 seconds.")
    
    filename = f"gemini_chat_{message.chat.id}.docx"
    
    try:
        extract_gemini_chat(url, filename)
        with open(filename, 'rb') as doc_file:
            bot.send_document(message.chat.id, doc_file)
        bot.delete_message(chat_id=message.chat.id, message_id=processing_msg.message_id)
            
    except Exception as e:
        bot.reply_to(message, f"❌ An error occurred:\n\n{str(e)}")
        
    finally:
        if os.path.exists(filename):
            os.remove(filename)

print("Bot is starting up...")
bot.polling(none_stop=True)
